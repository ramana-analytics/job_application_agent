import os
import re
import subprocess
from pathlib import Path

import pdfplumber
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file.
    
    Args:
        file_path: Path to the PDF file
    
    Returns:
        Extracted text from the PDF
    """
    text = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text.append(page.extract_text())
        return "\n".join(text)
    except Exception as e:
        raise ValueError(f"Error reading PDF: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
    
    Returns:
        Extracted text from the DOCX
    """
    try:
        doc = Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        return "\n".join(text)
    except Exception as e:
        raise ValueError(f"Error reading DOCX: {str(e)}")


def extract_text_from_latex(file_path: str) -> str:
    """Return the raw LaTeX source for a .tex file."""
    try:
        return Path(file_path).read_text(encoding="utf-8")
    except Exception as e:
        raise ValueError(f"Error reading LaTeX: {str(e)}")


def strip_latex_code_fences(source_text: str) -> str:
    """Remove a single outer Markdown code fence from generated LaTeX."""
    text = (source_text or "").strip()
    match = re.fullmatch(r"```(?:latex|tex)?\s*(.*?)\s*```", text, flags=re.IGNORECASE | re.DOTALL)
    return (match.group(1) if match else text).strip()


def _latex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    normalized = (
        text.replace("\u00a0", " ")
        .replace("\u2022", "*")
        .replace("\u2014", " -- ")
        .replace("\u2013", " - ")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
    )
    return "".join(replacements.get(ch, ch) for ch in normalized)


def plain_text_to_latex_source(text: str, document_title: str = "Resume") -> str:
    """Convert plain resume text into a minimal compilable LaTeX document."""
    lines = [line.rstrip() for line in (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    body: list[str] = []
    in_items = False

    def close_items() -> None:
        nonlocal in_items
        if in_items:
            body.append(r"\end{itemize}")
            body.append("")
            in_items = False

    for idx, raw_line in enumerate(lines):
        line = raw_line.strip()
        if not line:
            close_items()
            if body and body[-1] != "":
                body.append("")
            continue

        is_bullet = line.startswith(("- ", "* ", "• "))
        if is_bullet:
            if not in_items:
                body.append(r"\begin{itemize}")
                in_items = True
            body.append(rf"\item {_latex_escape(line[2:].strip())}")
            continue

        close_items()
        escaped = _latex_escape(line)
        is_heading = idx > 0 and line == line.upper() and 1 <= len(line.split()) <= 8
        if idx == 0:
            body.append(rf"{{\LARGE\bfseries {escaped}}}\par")
        elif idx == 1:
            body.append(rf"{escaped}\par")
        elif is_heading:
            body.append(rf"\section*{{{escaped}}}")
        else:
            body.append(rf"{escaped}\par")

    close_items()

    while body and body[-1] == "":
        body.pop()

    title = _latex_escape(document_title or "Resume")
    content = "\n".join(body) if body else rf"{title}\par"
    return "\n".join(
        [
            r"\documentclass[11pt]{article}",
            rf"\title{{{title}}}",
            r"\pagestyle{empty}",
            r"\setlength{\parindent}{0pt}",
            r"\begin{document}",
            content,
            r"\end{document}",
            "",
        ]
    )


def ensure_compilable_latex_source(source_text: str, document_title: str = "Resume") -> str:
    """Normalize Copilot/user LaTeX content into a complete compilable document."""
    cleaned = strip_latex_code_fences(source_text)
    if not cleaned:
        return plain_text_to_latex_source("", document_title)

    has_documentclass = bool(re.search(r"\\documentclass(?:\[[^\]]*\])?\{[^}]+\}", cleaned))
    has_begin = r"\begin{document}" in cleaned
    has_end = r"\end{document}" in cleaned
    has_command = bool(re.search(r"\\[a-zA-Z@]+", cleaned))

    if has_documentclass and has_begin and has_end:
        return cleaned.rstrip() + "\n"

    if has_command:
        body = re.sub(r"\\documentclass(?:\[[^\]]*\])?\{[^}]+\}", "", cleaned).strip()
        if r"\begin{document}" in body:
            body = body.split(r"\begin{document}", 1)[1]
        if r"\end{document}" in body:
            body = body.split(r"\end{document}", 1)[0]
        body = body.strip()
        return "\n".join(
            [
                r"\documentclass[11pt]{article}",
                r"\pagestyle{empty}",
                r"\setlength{\parindent}{0pt}",
                r"\begin{document}",
                body,
                r"\end{document}",
                "",
            ]
        )

    return plain_text_to_latex_source(cleaned, document_title)


def latex_source_to_plain_text(source_text: str) -> str:
    """Convert LaTeX source into plain text with Pandoc."""
    if not source_text.strip():
        return ""
    try:
        result = subprocess.run(
            ["pandoc", "-f", "latex", "-t", "plain"],
            input=source_text,
            text=True,
            capture_output=True,
            check=True,
        )
        return result.stdout.strip()
    except FileNotFoundError:
        raise ValueError("Pandoc is required to process LaTeX resumes")
    except subprocess.CalledProcessError as e:
        message = (e.stderr or e.stdout or str(e)).strip()
        raise ValueError(f"Error converting LaTeX to plain text: {message}")


def extract_text_from_latex_file(file_path: str) -> str:
    """Convert a .tex file into plain text for ATS/Copilot use."""
    return latex_source_to_plain_text(extract_text_from_latex(file_path))


def extract_resume_text(file_path: str) -> str:
    """
    Extract text from a resume (PDF, DOCX, TXT, or LaTeX source).
    
    Args:
        file_path: Path to the resume file
    
    Returns:
        Extracted resume text
    
    Raises:
        ValueError: If file format is not supported
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.txt':
        return Path(file_path).read_text(encoding="utf-8")
    elif ext == '.tex':
        return extract_text_from_latex(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def extract_resume_plain_text(file_path: str) -> str:
    """Extract the semantic/plain-text representation of a resume."""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    if ext == ".tex":
        return extract_text_from_latex_file(file_path)
    return extract_resume_text(file_path)


def save_resume_as_docx(content: str, output_path: str) -> bool:
    """
    Save resume content as a DOCX file.
    
    Args:
        content: The resume content (can be plain text or formatted)
        output_path: Path where the DOCX should be saved
    
    Returns:
        True if successful
    """
    try:
        doc = Document()

        # Apply resume-friendly page layout.
        section = doc.sections[0]
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

        # Set default body typography.
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)

        lines = [line.rstrip() for line in content.split("\n")]
        for idx, raw_line in enumerate(lines):
            line = raw_line.strip()
            if not line:
                continue

            # Detect likely name line at the top and emphasize it.
            if idx == 0 and len(line.split()) <= 4:
                para = doc.add_paragraph(line)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.runs[0]
                run.bold = True
                run.font.size = Pt(18)
                para.space_after = Pt(6)
                continue

            # Detect common section headers and style them.
            if line.isupper() or line.endswith(":"):
                para = doc.add_paragraph(line.replace(":", ""))
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.runs[0]
                run.bold = True
                run.font.size = Pt(12)
                para.space_before = Pt(8)
                para.space_after = Pt(3)
                continue

            # Convert bullet-like lines to native Word bullets.
            if line.startswith(("- ", "* ", "• ")):
                bullet_text = line[2:].strip()
                para = doc.add_paragraph(bullet_text, style="List Bullet")
                para.paragraph_format.space_after = Pt(2)
                continue

            para = doc.add_paragraph(line)
            para.paragraph_format.space_after = Pt(2)
        
        doc.save(output_path)
        return True
    except Exception as e:
        raise ValueError(f"Error saving DOCX: {str(e)}")


def save_resume_as_pdf(content: str, output_path: str) -> bool:
    """
    Save resume content as a PDF file.
    
    Args:
        content: The resume content
        output_path: Path where the PDF should be saved
    
    Returns:
        True if successful
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch
        
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom style for resume text
        custom_style = ParagraphStyle(
            'CustomResume',
            parent=styles['Normal'],
            fontSize=10,
            leading=12,
        )
        
        # Add paragraphs
        for line in content.split('\n'):
            if line.strip():
                story.append(Paragraph(line, custom_style))
                story.append(Spacer(1, 0.05*inch))
        
        doc.build(story)
        return True
    except Exception as e:
        raise ValueError(f"Error saving PDF: {str(e)}")


def convert_resume(input_path: str, output_path: str, output_format: str) -> bool:
    """
    Convert resume between formats (PDF <-> DOCX).
    
    Args:
        input_path: Path to input resume file
        output_path: Path for output resume file
        output_format: 'pdf' or 'docx'
    
    Returns:
        True if successful
    """
    _, ext = os.path.splitext(input_path)
    ext = ext.lower()

    if ext == '.tex':
        try:
            cmd = ["pandoc", input_path, "-o", output_path]
            if output_format.lower() == 'pdf':
                cmd.append("--pdf-engine=pdflatex")
            elif output_format.lower() != 'docx':
                raise ValueError(f"Unsupported output format: {output_format}")

            subprocess.run(cmd, capture_output=True, text=True, check=True)
            return True
        except FileNotFoundError:
            raise ValueError("Pandoc is required to convert LaTeX resumes")
        except subprocess.CalledProcessError as e:
            message = (e.stderr or e.stdout or str(e)).strip()
            raise ValueError(f"Error converting LaTeX resume: {message}")

    # Extract text from input
    content = extract_resume_text(input_path)

    # Save in output format
    if output_format.lower() == 'pdf':
        return save_resume_as_pdf(content, output_path)
    elif output_format.lower() == 'docx':
        return save_resume_as_docx(content, output_path)
    else:
        raise ValueError(f"Unsupported output format: {output_format}")
